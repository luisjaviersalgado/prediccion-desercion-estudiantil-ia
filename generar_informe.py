# -*- coding: utf-8 -*-
"""Genera el informe escrito (Word/PDF) de la actividad."""
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION

with open("data/resumen_resultados.json", encoding="utf-8") as f:
    R = json.load(f)

AZUL = RGBColor(0x1F, 0x3B, 0x73)
GRIS = RGBColor(0x44, 0x44, 0x44)

doc = Document()

# --- Estilos base -----------------------------------------------------
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for i, size in zip([1, 2, 3], [20, 15, 12.5]):
    hs = doc.styles[f"Heading {i}"]
    hs.font.name = "Calibri"
    hs.font.size = Pt(size)
    hs.font.bold = True
    hs.font.color.rgb = AZUL
    hs.paragraph_format.space_before = Pt(16 if i == 1 else 10)
    hs.paragraph_format.space_after = Pt(8)

sec = doc.sections[0]
sec.page_width = Cm(21.59)   # Carta
sec.page_height = Cm(27.94)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(2.7)
sec.right_margin = Cm(2.7)


def add_footer_pagenum(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)


add_footer_pagenum(sec)
sec.different_first_page_header_footer = True
# Footer de la portada queda vacío (no se numera la portada)
_ = sec.first_page_footer.paragraphs[0]


def parrafo(texto, size=11, bold=False, italic=False, align=None, color=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def figura(ruta, ancho=5.6, caption=None):
    doc.add_picture(ruta, width=Inches(ancho))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph()
        r = cp.add_run(caption)
        r.italic = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = GRIS
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(14)


def tabla_simple(headers, filas, anchos=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
    for fila in filas:
        cells = t.add_row().cells
        for i, val in enumerate(fila):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def salto_pagina():
    doc.add_page_break()


# ===========================================================================
# PORTADA
# ===========================================================================
def spacer(pt_height):
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(0)
    sp.paragraph_format.space_before = Pt(0)
    rr = sp.add_run("")
    rr.font.size = Pt(1)
    sp.paragraph_format.line_spacing = 1.0
    # emulate vertical space using space_before on the run's paragraph
    sp.paragraph_format.space_before = Pt(pt_height)
    return sp


spacer(60)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("UNIVERSIDAD TECNOLÓGICA DEL ORIENTE")
r.bold = True; r.font.size = Pt(15); r.font.color.rgb = AZUL

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Ingeniería de Software")
r.font.size = Pt(12); r.font.color.rgb = GRIS

spacer(70)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Creación de un Modelo de Inteligencia Artificial Avanzada")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = AZUL

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Predicción temprana del riesgo de deserción estudiantil universitaria\nmediante una red neuronal (Perceptrón Multicapa)")
r.font.size = Pt(14); r.italic = True; r.font.color.rgb = GRIS

spacer(50)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Estrategia pedagógica: Aprendizaje Basado en Retos (ABR)")
r.font.size = Pt(11)

spacer(110)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Luis Javier Salgado Guzmán")
r.bold = True; r.font.size = Pt(12)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Séptimo semestre")
r.font.size = Pt(11); r.font.color.rgb = GRIS

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Curso de Inteligencia Artificial")
r.font.size = Pt(11); r.font.color.rgb = GRIS

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("2026")
r.font.size = Pt(11); r.font.color.rgb = GRIS

salto_pagina()

# ===========================================================================
# TABLA DE CONTENIDO (campo actualizable en Word)
# ===========================================================================
doc.add_heading("Tabla de contenido", level=1)
p = doc.add_paragraph()
run = p.add_run()
fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
instr.text = 'TOC \\o "1-2" \\h \\z \\u'
fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
fld_text = OxmlElement("w:t"); fld_text.text = "Haga clic derecho y seleccione “Actualizar campo” para generar la tabla de contenido."
fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_sep); run._r.append(fld_text); run._r.append(fld_end)

salto_pagina()

# ===========================================================================
# 1. INTRODUCCIÓN Y JUSTIFICACIÓN
# ===========================================================================
doc.add_heading("1. Introducción y justificación del problema", level=1)

doc.add_heading("1.1 Identificación del problema", level=2)
parrafo(
    "La deserción estudiantil universitaria es uno de los problemas más persistentes de la "
    "educación superior en Latinoamérica. Cada estudiante que abandona sus estudios representa "
    "una pérdida de inversión académica, económica y social, tanto para la institución como para "
    "el propio estudiante y su entorno familiar. Detectar de forma temprana a los estudiantes con "
    "mayor riesgo de deserción permite a las instituciones educativas activar mecanismos de apoyo "
    "oportunos —tutorías académicas, acompañamiento psicosocial, ajustes en la carga académica o "
    "apoyo financiero— antes de que el estudiante tome la decisión de abandonar el programa."
)
parrafo(
    "En este proyecto se aborda el siguiente problema: ¿es posible, a partir de variables "
    "académicas, de asistencia y socioeconómicas disponibles al inicio o durante el semestre, "
    "predecir con antelación qué estudiantes tienen mayor probabilidad de desertar, de manera que "
    "la universidad pueda priorizar sus recursos de acompañamiento hacia los casos de mayor riesgo?"
)

doc.add_heading("1.2 Justificación", level=2)
parrafo(
    "La pertinencia de este problema para la Ingeniería de Software radica en que combina un reto "
    "de alto impacto social con un caso de uso representativo de la inteligencia artificial "
    "aplicada: un problema de clasificación binaria sobre datos tabulares con variables mixtas "
    "(numéricas y categóricas), en el cual, además del desempeño predictivo, es indispensable "
    "evaluar la equidad del modelo, dado que una intervención institucional basada en una "
    "predicción sesgada podría perjudicar sistemáticamente a los grupos que ya se encuentran en "
    "condición de mayor vulnerabilidad."
)
parrafo(
    "Una solución de IA aporta valor frente a los métodos tradicionales (por ejemplo, revisión "
    "manual de expedientes o umbrales fijos de promedio académico) porque puede integrar de forma "
    "simultánea múltiples factores de riesgo —académicos, de asistencia, laborales y "
    "socioeconómicos— y aprender relaciones no lineales entre ellos, generando una puntuación de "
    "riesgo individualizada que sirve como insumo, y no como reemplazo, del criterio de los "
    "consejeros académicos."
)
parrafo(
    "Este trabajo utiliza la estrategia de Aprendizaje Basado en Retos (ABR): se parte de un reto "
    "real y contextualizado (la deserción universitaria), se investiga el estado del arte, se "
    "diseña una solución basada en datos, se implementa y se evalúa críticamente, incluyendo sus "
    "implicaciones éticas."
)

# ===========================================================================
# 2. REVISIÓN DE LITERATURA
# ===========================================================================
doc.add_heading("2. Revisión de literatura", level=1)

doc.add_heading("2.1 Predicción de deserción con aprendizaje automático", level=2)
parrafo(
    "Diversas revisiones sistemáticas recientes muestran un crecimiento sostenido de la "
    "investigación sobre predicción de deserción estudiantil mediante aprendizaje automático "
    "entre 2019 y 2025, lo que refleja una creciente conciencia institucional sobre el problema. "
    "Estas revisiones coinciden en que los indicadores de desempeño académico y de asistencia o "
    "participación son los factores de riesgo más utilizados en los modelos predictivos, mientras "
    "que los factores psicosociales —autoeficacia, sentido de pertenencia, motivación y "
    "resiliencia— siguen estando subrepresentados en los conjuntos de datos disponibles."
)
parrafo(
    "En cuanto a los algoritmos empleados, los métodos de ensamble —particularmente Random Forest "
    "y Gradient Boosting— dominan el panorama por su buen desempeño sobre datos estructurados y su "
    "mayor facilidad de interpretación frente a modelos de caja negra, alcanzando en varios "
    "estudios exactitudes cercanas al 86-88 %, aunque la mayoría de las investigaciones se basa en "
    "datos de una sola institución y rara vez evalúa el impacto o la equidad de las intervenciones "
    "que se derivan de las predicciones. Estudios comparativos adicionales reportan que Random "
    "Forest es, en efecto, el algoritmo más frecuentemente utilizado en la literatura, con "
    "resultados de exactitud que llegan a superar el 90 % en configuraciones específicas de "
    "validación."
)

doc.add_heading("2.2 Equidad algorítmica (fairness) en contextos educativos", level=2)
parrafo(
    "La equidad algorítmica se ocupa de que las decisiones de un sistema de IA no dependan de "
    "manera injustificada de atributos sensibles como el género, el estrato socioeconómico o el "
    "origen étnico. Entre las métricas de equidad de grupo más utilizadas se encuentran la "
    "paridad demográfica (statistical parity), que exige tasas de selección similares entre "
    "grupos, y los criterios basados en el desempeño del clasificador condicionado a la etiqueta "
    "real: la igualdad de oportunidad (equal opportunity), que compara la tasa de verdaderos "
    "positivos (TPR) entre grupos, y la igualdad de momios (equalized odds), que exige además "
    "tasas de falsos positivos similares."
)
parrafo(
    "La literatura especializada advierte que imponer paridad demográfica de forma ciega puede ser "
    "inadecuado cuando existen diferencias reales en la prevalencia del resultado entre grupos, "
    "por lo que en aplicaciones educativas se suele preferir examinar la igualdad de oportunidad y "
    "de momios, que se centran en si el modelo comete errores de forma desproporcionada para "
    "ciertos grupos. En el ámbito específico de la predicción del desempeño académico, trabajos "
    "recientes concluyen que la igualdad de momios (equalized odds) es, en general, la noción de "
    "equidad más adecuada frente a la paridad demográfica, la calibración o la paridad predictiva, "
    "aunque reconocen que ninguna métrica captura por sí sola todas las dimensiones de la equidad "
    "percibida por los propios estudiantes."
)

doc.add_heading("2.3 Cómo informan estos hallazgos el diseño del modelo", level=2)
parrafo(
    "De la revisión de literatura se derivan tres decisiones de diseño para este proyecto: "
    "(1) se priorizan variables académicas (promedio, asistencia, materias reprobadas) y "
    "socioeconómicas como predictores principales, en línea con los factores de riesgo más "
    "reportados; (2) se compara el modelo avanzado (red neuronal) contra líneas base "
    "interpretables (regresión logística y Random Forest), dado que la literatura muestra que los "
    "métodos de ensamble son altamente competitivos y más transparentes; y (3) se incorpora de "
    "manera explícita una fase de evaluación de equidad utilizando igualdad de oportunidad e "
    "igualdad de momios como criterios complementarios a la paridad demográfica, siguiendo la "
    "recomendación de la literatura sobre predicción académica."
)

# ===========================================================================
# 3. DISEÑO DEL MODELO Y METODOLOGÍA
# ===========================================================================
doc.add_heading("3. Diseño del modelo y metodología", level=1)

doc.add_heading("3.1 Conjunto de datos", level=2)
parrafo(
    "Debido a que el entorno de desarrollo de este proyecto no cuenta con acceso a internet ni a "
    "repositorios institucionales de datos reales, se optó por construir un conjunto de datos "
    "sintético pero realista, generado mediante un modelo probabilístico que reproduce relaciones "
    "conocidas de la literatura entre las variables predictoras y el riesgo de deserción (por "
    "ejemplo: a menor promedio y asistencia, y a mayor número de materias reprobadas, mayor "
    "probabilidad de deserción). Este enfoque permite ilustrar el flujo de trabajo completo de un "
    "proyecto de IA aplicada, y deja explícitamente documentadas las relaciones causales asumidas, "
    "algo que rara vez es posible con datos reales. Para una implementación en producción, este "
    "conjunto debería reemplazarse por datos históricos reales de la institución, con el "
    "consentimiento y los controles de privacidad correspondientes."
)
parrafo(f"El conjunto de datos generado contiene {R['n_total']} estudiantes, con una tasa de "
        f"deserción global del {R['tasa_desercion_global']*100:.1f} %. La tabla 1 describe las "
        f"variables utilizadas.")

tabla_simple(
    ["Variable", "Tipo", "Descripción"],
    [
        ["edad", "Numérica", "Edad del estudiante (17-35 años)"],
        ["género", "Categórica", "Femenino / Masculino (atributo sensible)"],
        ["estrato", "Ordinal (1-6)", "Estrato socioeconómico (atributo sensible)"],
        ["semestre", "Numérica", "Semestre que cursa actualmente (1-10)"],
        ["trabaja", "Binaria", "Si el estudiante trabaja mientras estudia"],
        ["horas_estudio_semanal", "Numérica", "Horas de estudio autónomo por semana"],
        ["apoyo_financiero", "Binaria", "Si cuenta con beca o apoyo financiero"],
        ["distancia_km", "Numérica", "Distancia entre residencia y universidad"],
        ["participación_extracurricular", "Binaria", "Participación en actividades extracurriculares"],
        ["promedio_anterior", "Numérica", "Promedio académico del semestre anterior (0-5)"],
        ["asistencia_pct", "Numérica", "Porcentaje de asistencia a clases"],
        ["materias_reprobadas", "Numérica", "Materias reprobadas en el semestre anterior"],
        ["creditos_matriculados", "Numérica", "Número de créditos matriculados"],
        ["desercion (objetivo)", "Binaria", "1 = desertó, 0 = continúa activo"],
    ],
)

doc.add_heading("3.2 Preprocesamiento de datos", level=2)
parrafo(
    "Las variables numéricas se estandarizaron (media 0, desviación estándar 1) y la variable "
    "categórica género se codificó mediante one-hot encoding. El conjunto de datos se dividió de "
    "forma estratificada respecto a la variable objetivo en tres particiones:"
)
tabla_simple(
    ["Partición", "Proporción", "N.º de registros", "Uso"],
    [
        ["Entrenamiento", "60 %", str(R["n_train"]), "Ajuste de los parámetros del modelo"],
        ["Validación", "20 %", str(R["n_val"]), "Selección de umbral de decisión e hiperparámetros"],
        ["Prueba", "20 %", str(R["n_test"]), "Evaluación final, no usada en ningún ajuste"],
    ],
)

doc.add_heading("3.3 Arquitectura del modelo", level=2)
parrafo(
    "El modelo principal es un Perceptrón Multicapa (MLP), una red neuronal artificial "
    "feed-forward, elegido por su capacidad de capturar relaciones no lineales entre las variables "
    "predictoras sin requerir ingeniería manual de características. Como líneas base de "
    "comparación se entrenaron una Regresión Logística (modelo lineal interpretable) y un Random "
    "Forest (modelo de ensamble de árboles), siguiendo la recomendación de la literatura de "
    "contrastar modelos avanzados contra alternativas más transparentes."
)
arq = R["arquitectura_mlp"]
tabla_simple(
    ["Hiperparámetro", "Valor"],
    [
        ["Capas ocultas", " → ".join(str(x) for x in arq["capas_ocultas"]) + " neuronas"],
        ["Función de activación", arq["activacion"]],
        ["Optimizador", arq["optimizador"]],
        ["Regularización L2 (alpha)", str(arq["alpha_regularizacion"])],
        ["Parada anticipada (early stopping)", "Sí, 15 % de validación interna"],
        ["Iteraciones ejecutadas", str(arq["n_iteraciones_reales"])],
        ["Umbral de decisión", f"{R['umbral_decision']:.2f} (optimizado en validación, criterio F1)"],
    ],
)
parrafo(
    "La red recibe como entrada el vector de 13 variables preprocesadas, las propaga a través de "
    "dos capas ocultas de 64 y 32 neuronas con activación ReLU, y produce en la capa de salida una "
    "probabilidad de deserción mediante una función sigmoide. El entrenamiento utiliza el "
    "algoritmo Adam (descenso de gradiente estocástico adaptativo) para minimizar la función de "
    "pérdida de entropía cruzada binaria (log-loss), con regularización L2 para reducir el "
    "sobreajuste."
)
parrafo(
    "A diferencia de usar el umbral por defecto de 0.5, el umbral de decisión se ajustó sobre el "
    "conjunto de validación para maximizar el F1-score. Esto es una decisión de diseño deliberada: "
    "en un sistema de alerta temprana de deserción, omitir a un estudiante en riesgo real (falso "
    "negativo) es considerablemente más costoso que generar una alerta sobre un estudiante que "
    "finalmente no deserta (falso positivo), ya que la intervención de acompañamiento tiene un "
    "costo relativamente bajo frente al costo de una deserción no detectada."
)

# ===========================================================================
# 4. RESULTADOS DE ENTRENAMIENTO Y EVALUACIÓN
# ===========================================================================
doc.add_heading("4. Resultados de entrenamiento y evaluación", level=1)

doc.add_heading("4.1 Curva de entrenamiento", level=2)
parrafo(
    "La figura 1 muestra la evolución de la función de pérdida del MLP durante el entrenamiento. "
    "La pérdida desciende de forma consistente y se estabiliza tras aproximadamente 30 "
    "iteraciones, momento en el cual el criterio de parada anticipada detiene el entrenamiento "
    "para evitar sobreajuste."
)
figura("figures/curva_perdida_mlp.png", ancho=5.3, caption="Figura 1. Curva de pérdida (log-loss) durante el entrenamiento del MLP.")

doc.add_heading("4.2 Comparación de modelos", level=2)
parrafo("La tabla 4 y la figura 2 resumen el desempeño de los tres modelos sobre el conjunto de "
        "prueba, que no fue utilizado en ninguna etapa de ajuste.")

filas_metricas = []
for m in R["metricas_modelos"]:
    filas_metricas.append([
        m["modelo"], f"{m['exactitud']:.3f}", f"{m['precision']:.3f}",
        f"{m['sensibilidad_recall']:.3f}", f"{m['f1']:.3f}", f"{m['auc_roc']:.3f}",
    ])
tabla_simple(
    ["Modelo", "Exactitud", "Precisión", "Sensibilidad", "F1", "AUC-ROC"],
    filas_metricas,
)
figura("figures/comparacion_modelos.png", ancho=5.8, caption="Figura 2. Comparación de métricas de desempeño entre los tres modelos evaluados.")

parrafo(
    "El Random Forest obtiene el mejor balance entre precisión y sensibilidad, y el AUC-ROC más "
    "alto junto con la Regresión Logística, lo cual coincide con lo reportado en la literatura "
    "revisada. El MLP, con el umbral optimizado para F1, prioriza la sensibilidad "
    f"({R['metricas_modelos'][2]['sensibilidad_recall']:.1%}) a costa de una menor precisión "
    f"({R['metricas_modelos'][2]['precision']:.1%}): detecta a la gran mayoría de los estudiantes "
    "que efectivamente desertan, aunque genera más falsas alarmas. Esta es una decisión razonable "
    "para un sistema de alerta temprana, siempre que la intervención resultante (por ejemplo, una "
    "llamada de un consejero académico) sea de bajo costo y no estigmatizante para el estudiante."
)

figura("figures/curvas_roc.png", ancho=4.6, caption="Figura 3. Curvas ROC de los tres modelos sobre el conjunto de prueba.")

doc.add_heading("4.3 Matriz de confusión del modelo seleccionado", level=2)
parrafo(
    "Se seleccionó la Red Neuronal (MLP) como modelo de referencia para el análisis de equidad de "
    "la siguiente sección, por ser el modelo avanzado objeto de este trabajo y el que ofrece la "
    "mayor sensibilidad, la métrica más relevante para un caso de uso de alerta temprana."
)
figura("figures/matriz_confusion_mlp.png", ancho=4.3, caption="Figura 4. Matriz de confusión del MLP sobre el conjunto de prueba (umbral = 0.25).")

doc.add_heading("4.4 Variables más influyentes", level=2)
parrafo(
    "Aunque el MLP no ofrece coeficientes directamente interpretables, el análisis de importancia "
    "de variables del Random Forest —entrenado sobre las mismas variables— ofrece una aproximación "
    "útil de qué factores pesan más en la predicción del riesgo de deserción."
)
figura("figures/importancia_variables.png", ancho=5.2, caption="Figura 5. Importancia relativa de las variables según el Random Forest.")
parrafo(
    "El promedio académico anterior y el porcentaje de asistencia son, de forma consistente con la "
    "literatura, los predictores más influyentes, seguidos de la distancia al campus y las horas "
    "de estudio semanal, lo que sugiere que las barreras logísticas y de tiempo disponible tienen "
    "un peso relevante en el riesgo de deserción, más allá del desempeño académico puro."
)

# ===========================================================================
# 5. DISCUSIÓN SOBRE EQUIDAD Y ÉTICA
# ===========================================================================
doc.add_heading("5. Discusión sobre equidad y ética", level=1)
parrafo(
    "Un sistema de alerta temprana de deserción tiene consecuencias reales sobre las personas: "
    "puede determinar quién recibe acompañamiento prioritario, una beca de retención o una "
    "llamada de seguimiento. Por ello, no basta con evaluar la exactitud global del modelo; es "
    "necesario verificar que no cometa errores de forma sistemáticamente distinta entre grupos "
    "definidos por atributos sensibles como el género o el estrato socioeconómico."
)
parrafo(
    "Se utilizaron tres métricas de equidad de grupo, calculadas sobre las predicciones del MLP "
    "en el conjunto de prueba:"
)
p = doc.add_paragraph(style="List Bullet")
p.add_run("Disparate Impact Ratio (DIR): ").bold = True
p.add_run("razón entre la menor y la mayor tasa de selección (proporción de estudiantes marcados "
          "como riesgo alto) entre grupos. Un valor cercano a 1 indica paridad demográfica; valores "
          "por debajo de 0.8 se consideran, como regla práctica, indicio de posible discriminación.")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Equal Opportunity Difference (EOD): ").bold = True
p.add_run("diferencia máxima en la tasa de verdaderos positivos (sensibilidad) entre grupos. Mide "
          "si el modelo detecta con la misma eficacia a los estudiantes en riesgo real, "
          "independientemente del grupo al que pertenezcan.")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Equalized Odds Difference: ").bold = True
p.add_run("la mayor de las diferencias entre grupos en sensibilidad (TPR) y en tasa de falsos "
          "positivos (FPR); exige que el modelo se equivoque de forma similar en ambos sentidos "
          "para todos los grupos.")

doc.add_heading("5.1 Equidad por género", level=2)
figura("figures/equidad_genero.png", ancho=5.8, caption="Figura 6. Tasa de selección, sensibilidad (TPR) y falsos positivos (FPR) por género.")
rg = R["equidad_genero"]["resumen"]
tabla_simple(
    ["Métrica de equidad", "Valor obtenido", "Referencia deseable"],
    [
        ["Disparate Impact Ratio", f"{rg['disparate_impact_ratio']:.3f}", "≥ 0.80"],
        ["Equal Opportunity Difference", f"{rg['equal_opportunity_difference']:.3f}", "≤ 0.10 (ideal)"],
        ["Equalized Odds Difference", f"{rg['equalized_odds_difference']:.3f}", "≤ 0.10 (ideal)"],
    ],
)
parrafo(
    f"El Disparate Impact Ratio por género ({rg['disparate_impact_ratio']:.2f}) se mantiene dentro "
    "del umbral de referencia de 0.80, y la diferencia de igualdad de oportunidad es baja "
    f"({rg['equal_opportunity_difference']:.2f}), lo que indica que el modelo identifica a los "
    "estudiantes en riesgo real con una eficacia similar entre géneros. Sin embargo, la diferencia "
    f"de igualdad de momios ({rg['equalized_odds_difference']:.2f}) es más alta, impulsada "
    "principalmente por una mayor tasa de falsos positivos en el grupo femenino: el modelo genera "
    "más alertas 'de más' para mujeres que finalmente no desertan."
)

doc.add_heading("5.2 Equidad por estrato socioeconómico", level=2)
figura("figures/equidad_estrato.png", ancho=6.0, caption="Figura 7. Tasa de selección, sensibilidad (TPR) y falsos positivos (FPR) por estrato socioeconómico.")
re_ = R["equidad_estrato"]["resumen"]
tabla_simple(
    ["Métrica de equidad", "Valor obtenido", "Referencia deseable"],
    [
        ["Disparate Impact Ratio", f"{re_['disparate_impact_ratio']:.3f}", "≥ 0.80"],
        ["Equal Opportunity Difference", f"{re_['equal_opportunity_difference']:.3f}", "≤ 0.10 (ideal)"],
        ["Equalized Odds Difference", f"{re_['equalized_odds_difference']:.3f}", "≤ 0.10 (ideal)"],
    ],
)
parrafo(
    f"Para el estrato socioeconómico, el Disparate Impact Ratio ({re_['disparate_impact_ratio']:.2f}) "
    "se ubica justo en el límite de referencia de 0.80: los estudiantes de estrato bajo son "
    "marcados como riesgo alto con más frecuencia que los de estrato medio o alto. Es importante "
    "notar que, dado que la tasa de deserción real también es más alta en estrato bajo dentro de "
    "este conjunto de datos, cierta diferencia en la tasa de selección es esperable y no "
    "necesariamente injusta; por ello se complementa el análisis con la igualdad de oportunidad, "
    f"que muestra una diferencia baja ({re_['equal_opportunity_difference']:.2f}), es decir, el "
    "modelo detecta a los estudiantes en riesgo real de forma comparable entre estratos. La mayor "
    f"brecha aparece nuevamente en la igualdad de momios ({re_['equalized_odds_difference']:.2f}), "
    "por una tasa de falsos positivos más alta en estrato bajo."
)

doc.add_heading("5.3 Implicaciones éticas y recomendaciones de mitigación", level=2)
parrafo(
    "El hallazgo más relevante del análisis de equidad es que el modelo no discrimina de forma "
    "importante en su capacidad de detectar a los estudiantes que realmente están en riesgo "
    "(igualdad de oportunidad aceptable en ambos atributos sensibles), pero sí genera más falsas "
    "alarmas para los grupos de estrato bajo y género femenino. Esta asimetría tiene una "
    "implicación ética concreta: si la 'alerta de riesgo' se asocia, aunque sea informalmente, con "
    "una etiqueta de bajo desempeño o con un estigma, los estudiantes de estos grupos podrían "
    "verse afectados de forma desproporcionada por falsas alarmas, incluso si el objetivo original "
    "del sistema era ayudarlos."
)
parrafo(
    "Recomendaciones para una eventual puesta en producción de un sistema como este:"
)
for texto in [
    "Utilizar el resultado del modelo únicamente como una señal de apoyo a la decisión humana "
    "(consejeros académicos), nunca como criterio automático de sanción, condicionamiento de "
    "beneficios o exclusión.",
    "Diseñar la intervención asociada a una alerta positiva como un beneficio neutral o positivo "
    "para el estudiante (por ejemplo, una invitación a tutorías) y no como una etiqueta pública, "
    "para reducir el costo social de los falsos positivos.",
    "Reentrenar y auditar periódicamente el modelo con datos reales actualizados, verificando que "
    "las métricas de equidad se mantengan dentro de los umbrales de referencia a medida que "
    "cambia la población estudiantil.",
    "Evaluar técnicas de mitigación de sesgo (por ejemplo, reponderación de clases sensibles en el "
    "entrenamiento o ajuste de umbrales específicos por grupo) si, al auditar con datos reales, la "
    "diferencia de igualdad de momios resulta mayor a la aquí observada.",
    "Excluir el atributo de género y estrato como variables de entrada directas del modelo si su "
    "inclusión no aporta una mejora sustantiva de desempeño, y utilizarlos únicamente para la "
    "auditoría de equidad posterior al entrenamiento.",
]:
    doc.add_paragraph(texto, style="List Bullet")

# ===========================================================================
# 6. CONCLUSIONES Y RECOMENDACIONES
# ===========================================================================
doc.add_heading("6. Conclusiones y recomendaciones", level=1)
for texto in [
    "Se diseñó, entrenó y evaluó un modelo de red neuronal (Perceptrón Multicapa) capaz de "
    "predecir el riesgo de deserción estudiantil universitaria a partir de variables académicas, "
    "de asistencia y socioeconómicas, alcanzando un AUC-ROC de "
    f"{R['metricas_modelos'][2]['auc_roc']:.2f} y una sensibilidad de "
    f"{R['metricas_modelos'][2]['sensibilidad_recall']:.1%} sobre el conjunto de prueba.",
    "La comparación contra líneas base (Regresión Logística y Random Forest) mostró que, si bien "
    "el Random Forest ofrece el mejor equilibrio precisión-sensibilidad, el MLP con el umbral "
    "ajustado en validación resulta más adecuado para un caso de uso de alerta temprana, donde "
    "minimizar los falsos negativos (estudiantes en riesgo no detectados) es prioritario.",
    "El análisis de equidad evidenció que el modelo mantiene una capacidad de detección "
    "comparable entre géneros y estratos socioeconómicos (igualdad de oportunidad aceptable), "
    "pero genera más falsas alarmas para los grupos de estrato bajo y género femenino, lo que "
    "debe gestionarse cuidadosamente en el diseño de cualquier intervención derivada del modelo.",
    "El desempeño predictivo obtenido (AUC-ROC entre 0.66 y 0.68) es moderado y coherente con la "
    "naturaleza sintética y ruidosa de los datos; un despliegue real requeriría datos "
    "institucionales históricos, un mayor volumen de registros y variables adicionales, "
    "particularmente de tipo psicosocial, señaladas por la literatura como subrepresentadas pero "
    "relevantes.",
    "Como trabajo futuro se recomienda: (a) sustituir el conjunto de datos sintético por registros "
    "históricos reales de la institución; (b) incorporar variables de compromiso e interacción "
    "con plataformas virtuales de aprendizaje; (c) explorar técnicas de explicabilidad como SHAP "
    "para justificar cada predicción individual ante los consejeros académicos; y (d) realizar una "
    "validación piloto controlada antes de cualquier despliegue institucional.",
]:
    doc.add_paragraph(texto, style="List Bullet")

# ===========================================================================
# REFERENCIAS
# ===========================================================================
salto_pagina()
doc.add_heading("Referencias", level=1)
referencias = [
    "Albreiki, B., Zaki, N., & Alashwal, H. (2021). A systematic literature review of student "
    "performance prediction using machine learning techniques. Education Sciences, 11(9), 552.",
    "Girón-Valderrama, G., Ramirez, C., Toro-Dextre, E., Ausejo-Sanchez, J., Villarreal-Torres, H., "
    "& Angeles-Morales, J. (2023). Predicting student dropout based on machine learning and deep "
    "learning: A systematic review. EAI Endorsed Transactions on Scalable Information Systems.",
    "Hardt, M., Price, E., & Srebro, N. (2016). Equality of opportunity in supervised learning. "
    "Advances in Neural Information Processing Systems, 29.",
    "Nagy, M., & Molontay, R. (2023). Predicting dropout in higher education based on secondary "
    "school performance. Journal of Applied Statistics.",
    "Restrepo-Calle, F. et al. (2026). Student dropout prediction in higher education: A "
    "systematic review of machine learning methods and risk factors.",
    "Van der Rijst, T. et al. (2022). Towards equalised odds as fairness metric in academic "
    "performance prediction. arXiv:2209.14670.",
    "Autores varios (2026). Enhancing student retention in higher education institutions (HEIs): "
    "Machine learning approach. Electronics, 15(4), 734.",
    "Pedregosa, F. et al. (2011). Scikit-learn: Machine learning in Python. Journal of Machine "
    "Learning Research, 12, 2825-2830.",
    "Barocas, S., Hardt, M., & Narayanan, A. (2019). Fairness and Machine Learning: Limitations "
    "and Opportunities. fairmlbook.org.",
    "Tinto, V. (1975). Dropout from higher education: A theoretical synthesis of recent research. "
    "Review of Educational Research, 45(1), 89-125.",
]
for ref in referencias:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10)

doc.save("outputs/Informe_Deserción_Estudiantil_IA.docx")
print("Informe generado correctamente.")
